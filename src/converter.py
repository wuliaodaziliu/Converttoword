"""PDF转Word工具 - 转换核心模块"""

import os
import sys
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

def _runtime_base() -> Path:
    """返回运行时资源根目录。"""
    if getattr(sys, "frozen", False):
        return Path(getattr(sys, "_MEIPASS", Path(sys.executable).resolve().parent))
    return Path(__file__).resolve().parents[1]


def _poppler_candidates() -> list[Path]:
    """定位 bundled Poppler 的 bin 目录。"""
    candidates = [_runtime_base()]
    if getattr(sys, "frozen", False):
        candidates.append(Path(sys.executable).resolve().parent)
    return [base / "poppler" / "Library" / "bin" for base in candidates]


def _poppler_bin() -> Path | None:
    for path in _poppler_candidates():
        if (path / "pdfinfo.exe").exists() and (path / "pdftoppm.exe").exists():
            return path
    return None


_poppler_bin_path = _poppler_bin()
if _poppler_bin_path:
    os.environ["PATH"] = str(_poppler_bin_path) + os.pathsep + os.environ.get("PATH", "")

from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from docx.shared import Inches


class PDFConverter:
    """PDF转Word转换器"""

    def __init__(self, output_dir, dpi="Doc"):
        self.output_dir = Path(output_dir)
        self.dpi = {"Shot": 96, "Doc": 175, "Pic": 300}.get(dpi, 175)

    def convert(self, pdf_path, progress_callback=None):
        """转换单个PDF文件，返回 (success, message)"""
        pdf_path = Path(pdf_path)
        output_name = self._unique_filename(pdf_path.stem)
        output_path = self.output_dir / f"{output_name}.docx"

        try:
            poppler_path = str(_poppler_bin_path) if _poppler_bin_path else None
            if poppler_path is None:
                checked = "; ".join(str(path) for path in _poppler_candidates())
                return False, f"未找到内置 Poppler，请重新下载最新版 exe。已检查: {checked}"
            images = convert_from_path(
                str(pdf_path),
                dpi=self.dpi,
                fmt="png",
                poppler_path=poppler_path,
            )
            doc = Document()
            total = len(images)

            for i, img in enumerate(images, 1):
                img = self._resize_image(img)
                tmp_path = self._save_temp(img)

                try:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(tmp_path, width=Inches(6))
                    paragraph.alignment = 1  # CENTER
                finally:
                    if os.path.exists(tmp_path):
                        os.remove(tmp_path)

                if progress_callback:
                    progress_callback(i, total)

            doc.save(str(output_path))
            return True, str(output_path)

        except Exception as e:
            return False, str(e)

    def _resize_image(self, img):
        # 不再缩小图片，保持渲染后的原始尺寸，最大程度保留清晰度
        return img

    def _save_temp(self, img):
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        tmp.close()
        img.save(tmp.name, format='PNG')
        return tmp.name

    def _unique_filename(self, base):
        path = self.output_dir / f"{base}.docx"
        if not path.exists():
            return base
        i = 1
        while True:
            name = f"{base}_{i}"
            if not (self.output_dir / f"{name}.docx").exists():
                return name
            i += 1


def convert_multiple(pdf_files, output_dir, dpi="Doc", progress_callback=None):
    """并行转换多个PDF，返回 [(pdf_path, success, message), ...]"""
    results = []
    total_files = len(pdf_files)
    completed_files = 0

    def worker(pdf_file, idx):
        nonlocal completed_files

        def page_callback(page, total):
            overall = (completed_files + (page / total)) / total_files
            if progress_callback:
                progress_callback(page, total, overall)

        conv = PDFConverter(output_dir, dpi)
        success, msg = conv.convert(pdf_file, page_callback)
        completed_files += 1
        return pdf_file, success, msg

    with ThreadPoolExecutor(max_workers=4) as executor:
        future_to_file = {
            executor.submit(worker, f, i): f
            for i, f in enumerate(pdf_files)
        }
        for future in as_completed(future_to_file):
            pdf_file = future_to_file[future]
            try:
                pdf_file, success, msg = future.result()
                results.append((pdf_file, success, msg))
            except Exception as e:
                results.append((pdf_file, False, str(e)))

    return results
